from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Capacity_Agent_业务使用与指标说明.docx"

ACCENT = RGBColor(26, 82, 118)
LIGHT = "EAF2F8"
MID = "D6EAF8"
DARK = RGBColor(38, 50, 56)
MUTED = RGBColor(92, 105, 117)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color or DARK
    set_east_asia_font(r, "Microsoft YaHei")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name: str = "Microsoft YaHei", size: int = 10, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        set_east_asia_font(run, font_name)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def style_table(table, widths: list[float] | None = None, header_fill: str = LIGHT) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08
                set_paragraph_font(p, size=9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = ACCENT


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asia_font(run, "Microsoft YaHei")
        run.font.color.rgb = ACCENT if level <= 2 else DARK
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        for r in (r1, r2):
            set_east_asia_font(r, "Microsoft YaHei")
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        set_east_asia_font(r, "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_east_asia_font(r, "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = DARK


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    header_fill: str = LIGHT,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9, color=ACCENT)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, size=8 if len(value) > 60 else 9)
    style_table(table, widths, header_fill=header_fill)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 22), ("Subtitle", 11), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT if "Heading" in style_name or style_name == "Title" else MUTED

    header = section.header.paragraphs[0]
    header.text = "Capacity Agent | 业务使用与指标说明"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, size=8, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "内部评审与使用说明"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, size=8, color=MUTED)


def build_doc() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Capacity Agent 业务使用与指标说明")
    r.bold = True
    set_east_asia_font(r, "Microsoft YaHei")
    r.font.color.rgb = ACCENT
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("面向存储芯片制造厂 R6 主生产计划、WIP-aware 产能审查、投片承诺与周计划冻结")
    set_paragraph_font(subtitle, size=11, color=MUTED)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["适用对象", "生产计划、产能规划、制造运营、项目/产品组、系统开发与联调团队"],
            ["核心目标", "判断 R6 月度投片计划在现有产能和 WIP 约束下是否可执行，并在不可行时给出优化后的承诺建议"],
            ["使用边界", "当前适合做投片决策辅助、产能承诺评审和周计划冻结前校验；正式上线前需接入真实 MES/OEE/计划版本数据并完成接口级回归"],
            ["文档日期", "2026-05-05"],
        ],
        [3.4, 13.8],
        header_fill=MID,
    )

    add_heading(doc, "1. 项目解决的业务痛点", 1)
    add_body(doc, "Capacity Agent 是面向存储芯片制造场景的产能分析与计划优化工作台。它把 R6 月度投片需求、工艺路线、设备产能、OEE 日历和 WIP Lot 明细放到同一套分析链路中，帮助团队从“计划是否看起来可排”升级为“计划是否在真实 WIP 和瓶颈约束下可承诺”。")
    add_table(
        doc,
        ["业务痛点", "传统处理方式", "系统提供的能力"],
        [
            ["R6 投片计划缺少快速产能校验", "依赖 Excel 人工汇总，口径分散，结果滞后", "一键运行 RCCP，输出整体 Loading、瓶颈、缺口和可行性"],
            ["存储厂 WIP 周期长，忽略 WIP 容易误判", "只看新增投片需求，低估后续工序占机", "将 WIP 后续工序负载并入 RCCP 和生产计划主判定"],
            ["瓶颈判断靠经验，难解释", "按单个利用率或人工排序判断", "综合 Loading、Queue、Drift 形成瓶颈评分和主瓶颈"],
            ["不可行计划缺少优化方案", "人工反复削减产品需求", "LP/启发式优化给出调整后投片量、需求削减和瓶颈约束"],
            ["产能承诺和周计划冻结缺少状态口径", "可行/不可行与可冻结混用", "区分 capacity_feasible、commit_feasible、decision_readiness"],
            ["方案变更影响难量化", "单独拉表重算", "What-if 支持机台损失、需求增减等扰动对比"],
        ],
        [4.0, 5.2, 8.0],
    )

    add_heading(doc, "2. 典型使用场景", 1)
    add_bullets(
        doc,
        [
            "R6 月度主生产计划评审：导入每个产品的月度 wafer start 需求，判断当前产能和 WIP 约束下能否承诺。",
            "周计划冻结前校验：在冻结窗口前核对瓶颈 Loading、WIP 占机、需求削减和承诺可行性。",
            "产能承诺沟通：面向产品/销售/项目组解释哪些产品受限于哪些机台组，以及需要削减多少需求。",
            "瓶颈改善评估：通过 RCCP 热点、瓶颈评分和 What-if 分析判断加班、调机、外协或设备恢复的收益。",
            "Output 视角产出预测：从目标产出反推 WIP 是否足够，以及未来几周哪些产品能形成产出。",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "3. 使用前需要导入哪些资料", 1)
    add_body(doc, "前端支持导入一个 Excel 工作簿。当前要求 5 个必填 sheet，建议直接从页面下载最新模板后填充。字段名需要保持一致，系统会在导入时校验 sheet 和必填列。")
    add_table(
        doc,
        ["Sheet", "业务含义", "必填字段", "主要来源"],
        [
            ["route_master", "产品工艺路线和每道工序使用的机台组", "product_id, path_id, step_seq, tool_group_id, run_time_hr, batch_size", "MES/工艺路线主数据"],
            ["tool_groups", "机台组能力与设备数量", "tool_group_id, tool_group_name, area, n_machines, nameplate_throughput_wph", "设备资产系统/产能主数据"],
            ["oee", "设备有效产能日历", "fact_date, tool_group_id, availability, performance, quality, oee, available_hours", "OEE/设备绩效系统"],
            ["demand_plan", "R6 月度或周度投片需求", "time_window, product_id, wafer_count, plan_version, release_status", "主生产计划/R6 计划"],
            ["wip_lot_detail", "Lot 级 WIP 位置和状态", "lot_id, product_id, current_step_seq, wafer_count, percent_complete, lot_status", "MES WIP 明细"],
        ],
        [2.6, 4.0, 7.3, 3.3],
    )

    add_heading(doc, "3.1 推荐补充字段", 2)
    add_table(
        doc,
        ["Sheet", "推荐字段", "用途"],
        [
            ["route_master", "visit_count, route_version", "处理存储厂重复进站、多路线版本的工时口径"],
            ["demand_plan", "priority, contract_min, market_max, unit_profit, owner, approved_at", "支持 LP 优化、合约底线、市场上限和计划版本追溯"],
            ["wip_lot_detail", "good_wafer_count, wait_hours_so_far, remaining_wait_hours, hold_release_date, input_week", "提高 WIP 分摊、HOLD Lot 顺延、等待时间解释能力"],
            ["oee", "按日 fact_date 覆盖完整计划窗口", "R6 月度产能按目标月份日历汇总，避免用最新一天外推整月"],
        ],
        [3.0, 5.1, 9.1],
    )

    add_heading(doc, "3.2 数据准备质量要求", 2)
    add_bullets(
        doc,
        [
            "product_id 必须在 demand_plan、route_master、wip_lot_detail 中保持一致。",
            "tool_group_id 必须在 route_master、tool_groups、oee 中保持一致。",
            "R6 月度计划建议使用类似 R6-2026-05 或 2026-05 的 time_window，周计划可使用 2026-W17。",
            "available_hours 若已由设备系统计算，可直接作为有效小时；若为空或为 0，系统会基于 n_machines、availability、performance 估算。",
            "WIP Lot 若处于 HOLD，应提供 lot_status；若知道释放时间，建议填写 hold_release_date，避免把未来才释放的负载算入当前桶。",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "4. 前端使用流程", 1)
    add_table(
        doc,
        ["步骤", "页面操作", "系统动作", "产出结果"],
        [
            ["1", "进入全局控制台，选择内置样例或导入 Excel 工作簿", "解析数据集，校验 5 个必填 sheet", "生成数据集、产品数、机台组、路线行数、WIP Lot 数"],
            ["2", "选择需求时间窗", "读取 demand_plan 中该窗口的产品投片量", "展示 R6 月度投片需求合计"],
            ["3", "点击运行完整分析", "依次调用 RCCP、瓶颈分析、LP/DES、生产计划", "生成总览、热点、生产计划、冻结状态"],
            ["4", "进入生产计划页查看 LP 配置和输出", "根据目标、上下限、WIP 占机和产能矩阵调整计划", "输出承诺可行性、需求调整、产品级缺口"],
            ["5", "进入 Output/What-if 页做补充分析", "从产出目标或扰动条件重新计算", "输出产出缺口、预测完成度、扰动影响"],
            ["6", "必要时使用 Agent Chat", "用自然语言查询当前瓶颈和建议", "生成解释性问答结论"],
        ],
        [1.2, 5.0, 5.8, 5.2],
    )

    add_heading(doc, "5. 前端页面指标说明", 1)
    add_heading(doc, "5.1 全局控制台", 2)
    add_table(
        doc,
        ["指标/控件", "含义", "业务解读"],
        [
            ["当前数据集", "当前分析使用的数据来源，可为内置样例或 Excel 导入数据集", "所有页签共用同一个数据上下文"],
            ["需求时间窗", "demand_plan 中的 time_window", "决定本次分析对应 R6 月份或周计划窗口"],
            ["产品数", "数据集中参与计划或路线计算的产品数量", "用于判断本次计划覆盖范围"],
            ["机台组", "参与产能计算的 tool_group 数量", "用于判断设备约束覆盖范围"],
            ["路线行数", "route_master 的工序行数", "行数越多，路线颗粒度越细"],
            ["WIP Lot", "wip_lot_detail 中 Lot 数量", "决定 WIP-aware 分析是否有真实在制数据支撑"],
            ["R6 月度投片需求", "当前窗口下所有产品 wafer_count 合计", "本次主计划的投片需求总量"],
        ],
        [3.4, 6.0, 7.8],
    )

    add_heading(doc, "5.2 总览页", 2)
    add_table(
        doc,
        ["指标", "计算/来源", "如何解读"],
        [
            ["整体 Loading", "所有机台组需求小时 / 所有机台组可用小时", "越接近或超过 100%，整体产能越紧张；若含 WIP，会显示已含 WIP 后续负载"],
            ["可行性", "RCCP 主判定，任一机台组超过产能约束则不可行", "用于快速判断当前需求计划是否有明显产能缺口"],
            ["主瓶颈", "瓶颈评分最高的 tool_group", "优先用于产能改善、排产协调和异常跟踪"],
            ["DES 结论", "局部离散事件仿真结果", "用于验证排队效应是否会放大 RCCP 风险"],
            ["情景分类", "根据需求规模、瓶颈和约束识别场景", "提示推荐算法和求解器，帮助选择分析策略"],
        ],
        [3.2, 7.0, 7.0],
    )

    add_heading(doc, "5.3 RCCP 热点表", 2)
    add_table(
        doc,
        ["列名", "含义", "业务解释"],
        [
            ["Tool Group", "机台组编号", "定位瓶颈发生在哪类设备或区域"],
            ["新计划", "新投片需求带来的工时负载", "反映 R6 需求本身对该机台组的占用"],
            ["WIP", "当前 WIP 后续工序在该机台组的剩余占机小时", "存储厂长周期场景中非常关键，不能忽略"],
            ["Demand", "新计划 + WIP 的总需求小时", "用于和 Avail 比较判断是否过载"],
            ["Avail", "计划窗口内可用产能小时", "月度窗口按目标月份 OEE/available_hours 日历汇总"],
            ["Load", "Demand / Avail", "超过 100% 表示该机台组在当前计划下不可满足"],
            ["Gap", "Demand - Avail", "正数表示缺口小时，负数或 0 表示有余量"],
            ["Status", "healthy/warning/critical/overload", "用于快速识别风险级别"],
        ],
        [2.4, 6.3, 8.5],
    )

    add_heading(doc, "5.4 瓶颈评分", 2)
    add_table(
        doc,
        ["指标", "含义", "用途"],
        [
            ["loading_pct", "静态产能利用率", "判断机台是否接近或超过产能上限"],
            ["expected_wait_hours", "基于排队模型估算的等待小时", "判断高 Loading 是否会形成明显排队"],
            ["composite_score", "Loading、Queue、Drift 综合评分", "用于排序，分数越高越应优先处理"],
            ["severity", "风险等级", "用于区分 warning、critical、severe 等处理优先级"],
        ],
        [3.5, 5.8, 7.9],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "5.5 LP 与 DES", 2)
    add_table(
        doc,
        ["模块", "指标", "解释"],
        [
            ["LP Optimizer", "status", "optimal 表示最优求解完成；heuristic 表示启发式结果，正式冻结前仍需最优求解确认"],
            ["LP Optimizer", "objective_value", "优化目标值，当前页面以总利润或目标函数值展示"],
            ["LP Optimizer", "optimal_plan", "在产能约束下建议保留或调整后的产品投片量"],
            ["LP Optimizer", "binding_constraints", "已经打满的瓶颈机台组，是继续增产的主要约束"],
            ["DES Validation", "平均利用率", "仿真中机台实际占用比例，过高通常意味着排队风险"],
            ["DES Validation", "P95等待时间", "95% Lot 的等待时间上限，用于判断交期稳定性"],
            ["DES Validation", "周期时间", "仿真得到的平均处理周期，用于校验 RCCP 静态结论"],
        ],
        [3.2, 4.0, 10.0],
    )

    add_heading(doc, "5.6 生产计划页", 2)
    add_table(
        doc,
        ["指标", "含义", "决策口径"],
        [
            ["冻结状态", "decision_readiness 的业务化展示", "可冻结、可承诺、待求解器确认、不可冻结"],
            ["承诺可行性", "commit_feasible", "只有最优求解确认或原计划无需调整且可行时，才可作为正式承诺"],
            ["整体 Loading", "生产计划口径下的平均 Loading", "判断调整后计划整体紧张程度"],
            ["需求调整", "原始需求总量 - 调整后需求总量", "表示为了满足产能约束需要削减的 wafer 数"],
            ["原计划结论", "未考虑 WIP 或优化前的可行性", "用于对比原计划是否天然可执行"],
            ["WIP校正后", "加入 WIP 后续负载后的产能可行性", "更贴近存储厂真实执行约束"],
            ["WIP占机", "WIP 后续工序累计占机小时", "数值越大，新投片可用空间越小"],
            ["WIP影响", "WIP 校正后 Loading - 原计划 Loading", "衡量 WIP 对计划可执行性的压力增量"],
        ],
        [3.2, 6.4, 7.6],
    )

    add_heading(doc, "5.7 产品计划对比表", 2)
    add_table(
        doc,
        ["列名", "含义", "使用建议"],
        [
            ["产品", "product_id", "对应 R6 计划中的产品型号"],
            ["目标", "原始计划 wafer_count", "业务希望投片或交付的目标量"],
            ["原可达", "未考虑 WIP 校正时估算可达量", "用于说明传统口径可能高估的部分"],
            ["WIP校正后", "并入 WIP 后可支持的量", "作为更保守的计划执行参考"],
            ["原缺口", "目标 - 原可达", "判断原计划自身缺口"],
            ["校正后缺口", "目标 - WIP校正后可达", "判断真实执行缺口"],
            ["优先级", "priority", "用于计划削减或保供排序"],
            ["利润", "unit_profit × 可分配量", "用于最大利润目标下的产品排序"],
        ],
        [3.0, 5.4, 8.8],
    )

    add_heading(doc, "5.8 Output / What-if / Agent", 2)
    add_table(
        doc,
        ["页面区域", "指标/功能", "解释"],
        [
            ["Output 视角", "产出目标周", "以某一周或窗口的目标产出作为反推起点"],
            ["Output 视角", "WIP占比", "WIP 后续工序小时占总需求小时的比例"],
            ["Output 视角", "产出缺口", "预测产出低于目标的 wafer 数"],
            ["Output 视角", "预测完成度", "预测产出 / 目标产出，用于判断产出目标达成概率"],
            ["Output 预测", "总 WIP、预测周数、产品数", "说明预测覆盖多少在制品、多少周、多少产品"],
            ["What-if", "机台损失/需求变化", "模拟设备不可用或产品需求变化对 Loading 和可行性的影响"],
            ["Agent Chat", "自然语言问答", "用于让业务用户快速询问瓶颈、风险、建议和配置解释"],
        ],
        [3.2, 4.6, 9.4],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. 决策口径与结果解释", 1)
    add_table(
        doc,
        ["状态/字段", "含义", "是否可用于正式投片承诺"],
        [
            ["commit_ready_optimal / 可冻结", "计划经最优求解器验证，且产能约束满足", "是"],
            ["commit_ready_no_adjustment / 可承诺", "原计划无需调整即满足产能约束", "通常可以，但仍建议结合业务审批"],
            ["solver_required_for_final_commit / 待求解器确认", "启发式或非最优结果显示产能可行，但未完成正式最优求解", "不建议直接冻结"],
            ["not_commit_ready / 不可冻结", "计划仍超产能或关键约束不可满足", "否"],
            ["capacity_feasible", "仅表示调整后在产能小时上可行", "不能单独等同于正式承诺"],
            ["commit_feasible", "系统用于正式承诺展示的主字段", "是，以该字段作为前端承诺可行性依据"],
        ],
        [5.0, 8.6, 3.6],
    )

    add_heading(doc, "7. 建议的评审使用方式", 1)
    add_bullets(
        doc,
        [
            "先看总览：确认整体 Loading、可行性、主瓶颈和 DES 结论。",
            "再看 RCCP 热点：定位哪些机台组是过载来源，区分新计划负载和 WIP 负载。",
            "再看生产计划：以冻结状态和承诺可行性作为是否可进入投片承诺的主判断。",
            "若显示待求解器确认：不要直接冻结，应补装/启用正式 LP 求解器并重新运行。",
            "若显示不可冻结：结合产品优先级、contract_min、market_max 和瓶颈机台改善方案做下一轮计划调整。",
            "评审留档时建议记录数据集版本、time_window、plan_version、release_status、运行时间和主要瓶颈。",
        ],
    )

    add_heading(doc, "8. 注意事项与上线前补强", 1)
    add_table(
        doc,
        ["事项", "说明"],
        [
            ["数据源", "当前支持 Excel 导入和内置样例；正式上线建议与 MES、OEE、MPS/R6 系统对接，减少人工复制错误"],
            ["求解器", "若未安装 Pyomo/HiGHS 等最优求解器，系统可能返回 heuristic，此时只能做参考，不能直接冻结"],
            ["WIP 精度", "WIP 分摊依赖 current_step_seq、lot_status、剩余等待时间和路线工时，字段越完整结果越可信"],
            ["产能日历", "R6 月度产能应覆盖目标月份所有有效日期，避免只导入某一天 OEE 导致月度产能偏差"],
            ["流程闭环", "正式用于承诺前，应增加审批流、数据版本锁定、结果导出和审计追溯"],
        ],
        [4.0, 13.2],
    )

    add_heading(doc, "9. 一句话总结", 1)
    add_body(doc, "Capacity Agent 解决的是“R6 投片计划在真实产能与 WIP 约束下到底能不能承诺”的问题。业务上，它把计划、路线、设备、OEE 和 WIP 拉到统一口径；系统上，它把 RCCP、瓶颈识别、LP 优化、DES 校验、Output 预测和 What-if 推演整合到同一个前端工作台，帮助项目组把计划评审从经验判断推进到可解释、可追溯、可优化的决策流程。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
