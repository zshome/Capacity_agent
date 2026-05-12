from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Capacity_Agent_Excel模板字段审查说明.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="0B2545")
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_callout(doc, label, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.font.size = Pt(10.5)
    p.add_run("  " + body)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


CORE_ROWS = [
    ("A 类", "算法必需字段", "导入时必须校验，缺失会导致产能计算、WIP 负荷或投片优化失真。"),
    ("B 类", "正式决策增强字段", "建议保留，用于 R6 承诺、冻结、优先级、复杂 Path、备机和场景分析。"),
    ("C 类", "展示 / 治理 / 未来扩展字段", "可选字段，不建议强制填写；用于解释、追溯、页面展示或后续算法扩展。"),
]

SHEETS = [
    (
        "route_master",
        "产品工艺路径与设备组负荷折算表",
        [
            ("product_id", "产品与工艺路径匹配", "A 必填", "核心主键，缺失无法计算该产品产能需求。"),
            ("path_id", "区分主路径、备选路径或版本路径", "A 必填", "正式使用时建议与 route_version 配合。"),
            ("step_seq", "识别工艺顺序与 WIP 所在站点", "A 必填", "WIP 只应计算当前站点之后的剩余负荷。"),
            ("tool_group_id", "映射瓶颈设备组", "A 必填", "RCCP 产能约束的资源维度。"),
            ("run_time_hr", "单步加工时间", "A 必填", "用于折算小时负荷。"),
            ("batch_size", "批量大小", "A 必填", "用于批量折算小时 / 片。"),
            ("visit_count", "重复进站次数", "B 建议保留", "支持存储芯片制造中的重入路径。"),
            ("route_version", "路径版本", "B 建议保留", "支持 current、R6、scenario 等版本治理。"),
            ("step_name", "工序名称展示", "C 可选", "主要用于前端解释与业务沟通。"),
        ],
    ),
    (
        "tool_groups",
        "设备组与产能资源主数据",
        [
            ("tool_group_id", "产能资源主键", "A 必填", "连接 Route、OEE、WIP 与产能结果。"),
            ("tool_group_name", "设备组名称", "B 建议保留", "便于页面与汇报展示。"),
            ("area", "工艺区域", "B 建议保留", "支持按区域查看产能与瓶颈。"),
            ("n_machines", "设备台数", "A 必填", "available_hours 缺失时用于折算设备可用小时。"),
            ("nameplate_throughput_wph", "理论吞吐", "B 建议保留", "用于瓶颈、排队与 DES 分析解释。"),
            ("process_type", "工艺分类", "C 可选", "当前偏展示与治理。"),
            ("is_active", "资源是否启用", "B 建议保留", "建议后续算法真正过滤未启用资源。"),
        ],
    ),
    (
        "oee",
        "设备组可用产能与效率数据",
        [
            ("fact_date", "产能日期 / 周期", "A 必填", "用于匹配计划窗口与历史趋势。"),
            ("tool_group_id", "关联设备组", "A 必填", "连接资源主数据。"),
            ("available_hours", "实际可用小时", "A 必填", "最直接的产能约束输入。"),
            ("availability", "可用率", "B 建议保留", "可作为 available_hours 的补充或解释。"),
            ("performance", "性能效率", "B 建议保留", "可作为 OEE 与产能折算依据。"),
            ("quality", "质量效率 / 良率", "C 可选", "若已提供 oee 和 available_hours，可作为解释字段。"),
            ("oee", "综合设备效率", "B 建议保留", "用于趋势、看板和效率解释。"),
        ],
    ),
    (
        "demand_plan",
        "R6 月度投片需求与计划治理数据",
        [
            ("time_window", "计划窗口", "A 必填", "用于月度、周度或冻结周期匹配。"),
            ("product_id", "计划产品", "A 必填", "连接路径和 WIP。"),
            ("wafer_count", "需求投片量", "A 必填", "主生产计划输入。"),
            ("priority", "计划优先级", "B 建议保留", "不可行时用于排序与取舍。"),
            ("contract_min", "最低承诺量", "B 正式决策建议保留", "支撑产能承诺与客户保障。"),
            ("market_max", "最大市场需求", "B 建议保留", "作为优化上限，避免过量投片。"),
            ("unit_profit", "单位收益", "B 视目标保留", "用于利润最大化场景。"),
            ("plan_version", "计划版本", "B 建议保留", "支撑 R6、冻结版、模拟版追溯。"),
            ("release_status", "发布状态", "B 建议保留", "区分草稿、冻结、发布计划。"),
            ("owner", "责任人", "C 可选", "治理与追溯字段。"),
            ("approved_at", "审批时间", "C 可选", "治理与审计字段。"),
        ],
    ),
    (
        "wip_lot_detail",
        "在制品批次明细",
        [
            ("lot_id", "WIP 批次识别", "A 必填", "用于批次追踪与去重。"),
            ("product_id", "WIP 所属产品", "A 必填", "连接产品路径。"),
            ("current_step_seq", "当前工序位置", "A 必填", "决定剩余工艺负荷。"),
            ("wafer_count", "WIP 数量", "A 必填", "用于剩余负荷与出货口径计算。"),
            ("percent_complete", "完成比例", "A 必填", "用于估算剩余负荷。"),
            ("lot_status", "批次状态", "A 必填", "RUN、WAIT、HOLD 等状态影响时间偏移。"),
            ("good_wafer_count", "良片数量", "B 建议保留", "用于输出与良率口径。"),
            ("wait_hours_so_far", "已等待时间", "B 建议保留", "用于排队与滞留分析。"),
            ("remaining_wait_hours", "剩余等待时间", "B 建议保留", "用于 WIP 负荷时间偏移。"),
            ("hold_release_date", "HOLD 释放日期", "B 建议保留", "正式排程场景需要。"),
            ("input_week", "入站周", "C 可选", "用于追溯和周期分析。"),
        ],
    ),
    (
        "tool_master",
        "单台设备主数据，复杂 Path 场景使用",
        [
            ("tool_id", "单机主键", "A 复杂 Path 必填", "用于单机能力与备机关系。"),
            ("tool_group_id", "所属设备组", "A 复杂 Path 必填", "连接单机与设备组。"),
            ("tool_name", "设备名称", "C 可选", "主要用于展示。"),
            ("status", "设备状态", "B 建议保留", "应进一步用于过滤 DOWN / PM 设备产能。"),
            ("uptime", "单机可用率", "A 复杂 Path 必填", "用于计算单机可用小时。"),
            ("loss_time", "损失时间", "A 复杂 Path 必填", "用于扣减实际可用产能。"),
        ],
    ),
    (
        "process_master",
        "工序主数据，复杂 Path 场景使用",
        [
            ("process_id", "工序主键", "A 复杂 Path 必填", "连接能力矩阵与备机规则。"),
            ("process_name", "工序名称", "C 可选", "用于展示和业务解释。"),
            ("process_seq", "工序顺序", "A 复杂 Path 必填", "用于排序与路径判断。"),
            ("area", "工艺区域", "C 可选", "用于区域展示。"),
        ],
    ),
    (
        "tool_process_capability",
        "产品-工序-设备能力矩阵",
        [
            ("product_id", "产品维度能力", "A 复杂 Path 必填", "判断某产品能否在某设备加工。"),
            ("process_id", "工序维度能力", "A 复杂 Path 必填", "连接工序。"),
            ("tool_id", "设备维度能力", "A 复杂 Path 必填", "连接单机能力。"),
            ("can_run", "是否可加工", "A 复杂 Path 必填", "决定可行 Path。"),
            ("run_time_hr", "加工时间", "A 复杂 Path 必填", "用于单机负荷折算。"),
            ("batch_size", "批量大小", "A 复杂 Path 必填", "用于批量折算。"),
            ("path_type", "路径类型", "C 可选", "当前主要用于展示与治理。"),
            ("visit_count", "重复进站次数", "B 建议保留", "支持存储制造重入路径。"),
        ],
    ),
    (
        "backup_path",
        "备机与替代路径规则",
        [
            ("primary_tool_id", "主设备", "A 复杂 Path 必填", "定义被替代设备。"),
            ("backup_tool_id", "备机设备", "A 复杂 Path 必填", "定义可替代设备。"),
            ("process_id", "适用工序", "A 复杂 Path 必填", "备机通常按工序限定。"),
            ("enable_rule", "启用规则", "B 建议保留", "建议后续算法解释规则，如主机满载、主机停机等。"),
            ("switch_cost", "切换成本", "C 当前可选", "当前可作为未来优化目标或惩罚项。"),
        ],
    ),
]


def main():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Capacity Agent Excel 模板字段审查说明")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("用于需求评审留档、前端导入校验与后续算法优化")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_callout(
        doc,
        "审查结论",
        "当前模板已经覆盖 R6 月度投片、产能约束、WIP、复杂 Path、备机能力等正式投片决策所需核心数据。后续不建议继续简单增加字段，而应按“标准版 / 正式决策版”分层管理，并让算法真正消费关键治理字段。",
        fill="EEF6FF",
    )

    doc.add_heading("1. 字段分级原则", level=1)
    add_table(doc, ["类型", "含义", "处理建议"], CORE_ROWS, [0.9, 1.7, 4.7])

    doc.add_heading("2. 推荐模板分层", level=1)
    add_bullet(doc, "标准版模板：保留 route_master、tool_groups、oee、demand_plan、wip_lot_detail，覆盖常规 RCCP、WIP 口径分析和计划可行性判断。")
    add_bullet(doc, "正式决策版模板：在标准版基础上增加 tool_master、process_master、tool_process_capability、backup_path，覆盖复杂 Path、单机能力、备机切换和正式投片承诺。")
    add_bullet(doc, "README 与 field_dictionary 不参与算法计算，但建议保留，用于业务方填报说明、字段口径统一和模板自解释。")

    doc.add_heading("3. 字段用途与必要性审查", level=1)
    for sheet, desc, rows in SHEETS:
        doc.add_heading(f"{sheet}：{desc}", level=2)
        add_table(doc, ["字段", "主要用途", "必要性", "审查说明"], rows, [1.45, 1.95, 1.35, 2.55])

    doc.add_heading("4. 当前模板与算法匹配问题", level=1)
    add_bullet(doc, "demand_plan 中 plan_version、release_status 当前更偏治理字段，但若目标是 R6 正式投片承诺，应继续保留。")
    add_bullet(doc, "tool_master.status 当前建议从展示字段升级为算法字段，用于扣除 DOWN、PM 或不可用设备产能。")
    add_bullet(doc, "backup_path.enable_rule 与 switch_cost 当前更像预留字段，后续应进入优化器，否则前端应标注为“可选 / 暂不参与计算”。")
    add_bullet(doc, "oee 中 available_hours 是产能约束的优先字段；availability、performance、quality、oee 应作为补充解释或缺失时的折算依据。")
    add_bullet(doc, "route_master 与 tool_process_capability 不建议直接合并，前者面向设备组级 RCCP，后者面向单机级复杂 Path 和备机能力。")

    doc.add_heading("5. 后续优化建议", level=1)
    add_bullet(doc, "前端导入区按标准版 / 正式决策版展示必填 Sheet，降低业务方首次导入压力。")
    add_bullet(doc, "字段字典中增加“是否参与当前算法计算”列，避免业务误以为所有字段都会影响结果。")
    add_bullet(doc, "导入校验区给出缺字段、空值、类型错误、跨 Sheet 主键不匹配等可操作提示。")
    add_bullet(doc, "算法侧补齐 is_active、tool status、backup enable_rule、switch_cost 的消费逻辑，确保模板字段与正式决策能力一致。")
    add_bullet(doc, "保留复杂 Path 四张表的整体一致性校验：只要导入其中一张，就要求四张均完整导入，避免半套复杂路径数据导致误判。")

    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Capacity Agent | Excel 模板字段审查说明")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
