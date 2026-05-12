from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "公司产能计算方法解读与项目优化建议.docx"

ACCENT = RGBColor(31, 96, 156)
GREEN = RGBColor(46, 125, 50)
ORANGE = RGBColor(209, 126, 26)
DARK = RGBColor(28, 42, 58)
MUTED = RGBColor(91, 105, 123)
LIGHT_BLUE = "EAF3FF"
LIGHT_GREEN = "EAF7EA"
LIGHT_ORANGE = "FFF3E0"


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color or DARK
    set_east_asia_font(r)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths: list[float] | None = None, header_fill: str = LIGHT_BLUE) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = ACCENT


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    header_fill: str = LIGHT_BLUE,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=ACCENT, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v, size=8 if len(v) > 70 else 9)
    style_table(table, widths, header_fill)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_east_asia_font(run)
        run.font.color.rgb = ACCENT if level <= 2 else DARK


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.16
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    set_east_asia_font(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
        set_east_asia_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_ORANGE, color: RGBColor = ORANGE) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title + "：")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = color
    set_east_asia_font(r1)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    set_east_asia_font(r2)
    style_table(table, [17.2], header_fill=fill)
    doc.add_paragraph()


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Title"].font.size = Pt(22)
    doc.styles["Title"].font.color.rgb = ACCENT
    doc.styles["Subtitle"].font.size = Pt(11)
    doc.styles["Subtitle"].font.color.rgb = MUTED

    header = section.header.paragraphs[0]
    header.text = "Capacity Agent | 公司产能计算方法解读"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "内部评审材料"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    r = title.add_run("公司产能计算方法解读与 Capacity Agent 优化建议")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    set_east_asia_font(r)

    subtitle = doc.add_paragraph(style="Subtitle")
    r = subtitle.add_run("基于《产能模拟计算_诉求&情景_0424_晋华.xlsx》的业务方法拆解")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    set_east_asia_font(r)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["文档目的", "解读公司现有产能计算诉求、情景分类和复杂 path 含义，并提出 Capacity Agent 后续优化方向。"],
            ["源文件", "产能模拟计算_诉求&情景_0424_晋华.xlsx"],
            ["源文件结构", "单 Sheet，37 行 × 13 列，107 个非空单元格；未发现公式、命名区域、图表或数据表。"],
            ["核心结论", "Excel 本身不是公式模型，而是产能计算业务规则和情景分类说明；复杂度集中在机台 path 差异、backup、可行性矩阵和分配优化。"],
        ],
        [3.4, 13.8],
        header_fill=LIGHT_GREEN,
    )

    add_heading(doc, "1. Excel 文件表达的核心诉求", 1)
    add_body(
        doc,
        "源文件开头明确提出两个核心诉求：一是产能计算准确率高，二是系统可以给出产能提升的最佳或最快路径建议。这说明该方法不是只要求算出一个静态产能值，而是希望在复杂设备约束下判断瓶颈，并给出改善方向。",
    )
    add_callout(
        doc,
        "标准产能公式",
        "24hr / TTL_TC × batch_size × 30天 × (uptime - loss_time)。该公式适合 path 简单、机台能力一致或可汇总的情景。",
        fill=LIGHT_BLUE,
        color=ACCENT,
    )
    add_body(
        doc,
        "源文件还特别注明：backup 情况下要求产能极大化。也就是说，当存在备用路径或替代机台时，计算目标不应只是按默认路径计算产能，而是要在所有可用路径中寻找最大产出组合。",
    )

    add_heading(doc, "2. 五类产能计算情景", 1)
    add_table(
        doc,
        ["情景", "产品", "机台", "Path 状况", "Backup", "建议计算方式"],
        [
            ["情景1", "多产品", "单台", "相同", "有或无", "标准产能计算"],
            ["情景2", "多产品", "多台", "相同", "无", "标准产能计算"],
            ["情景3", "多产品", "多台", "相同", "有", "标准产能计算，backup 可作为能力补充"],
            ["情景4", "多产品", "多台", "不同", "无", "穷举、找约束、分配模型"],
            ["情景5", "多产品", "多台", "不同", "有", "穷举、找约束、分配模型，并以产能极大化为目标"],
        ],
        [2.2, 2.4, 2.4, 3.0, 2.4, 5.4],
    )
    add_callout(
        doc,
        "关键分界点",
        "当机台 path 相同时，可以把机台组能力汇总后套标准公式；当机台 path 不同时，不能简单汇总，需要根据产品、制程、机台可行性做分配优化。",
    )

    add_heading(doc, "3. 什么是复杂 path", 1)
    add_body(
        doc,
        "这里的复杂 path 指：同一产品或制程在不同机台上的可加工路径不完全一致，导致产能不能简单按机台数量相加。复杂 path 的本质不是“有几台机”，而是“哪台机能跑哪个产品的哪个制程，以及怎么分配才不会卡住”。",
    )
    add_table(
        doc,
        ["机台", "制程1", "制程2", "制程3", "业务含义"],
        [
            ["机台1", "可跑", "可跑", "不可跑", "可承担前段工序，但不能支撑制程3"],
            ["机台2", "可跑", "不可跑", "可跑", "可能成为制程3的唯一或主要瓶颈"],
            ["机台3", "不可跑", "可跑", "不可跑", "只能补充制程2产能"],
        ],
        [2.3, 2.2, 2.2, 2.2, 8.3],
        header_fill=LIGHT_ORANGE,
    )
    add_body(
        doc,
        "如果产品必须依次经过制程1、制程2、制程3，那么制程3只能由机台2承担。即使总机台数很多，只要关键制程可用机台少，瓶颈仍然会集中在少数机台上。因此复杂 path 场景必须做机台-制程-产品分配，而不是只看总产能。",
    )

    add_heading(doc, "4. 后半部分矩阵的业务含义", 1)
    add_body(
        doc,
        "Excel 后半部分用两个产品、三个机台、三个制程举例。产品一需求为 1000 PCS，产品二需求为 500 PCS。矩阵中的 V 表示该机台可跑该制程，X 或 V→X 表示能力变化或不可用，V(可以) 可能表示 backup 或替代路径可行。",
    )
    add_table(
        doc,
        ["符号", "解释", "对算法的影响"],
        [
            ["V", "机台可以加工该制程", "可作为分配模型中的可行边"],
            ["X", "机台不能加工该制程", "该组合不可分配"],
            ["V→X", "原本可加工但当前变为不可用", "需要按当前状态剔除或作为异常情景"],
            ["V(可以)", "可能表示备用能力或条件满足时可用", "需要区分 primary path 与 backup path，并设置启用条件"],
        ],
        [2.2, 5.2, 9.8],
    )
    add_callout(
        doc,
        "解读",
        "这部分矩阵实际是在定义产品-制程-机台可行性矩阵。它是复杂 path 计算的基础输入，也是后续 LP/MIP、最大流或分配模型的核心约束。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "5. 对 Capacity Agent 当前项目的影响", 1)
    add_body(
        doc,
        "当前 Capacity Agent 已经具备 R6 计划导入、WIP-aware RCCP、瓶颈识别、LP 优化、DES 校验和冻结状态判定能力。若要完全贴合该 Excel 所表达的公司方法，需要进一步从 tool_group 级产能模型扩展到 tool_id 级、制程级、path 级分配模型。",
    )
    add_table(
        doc,
        ["公司方法要求", "当前项目已有能力", "需要补强"],
        [
            ["标准产能公式", "RCCP 可计算产品-机台组工时与 Loading", "补充 TTL_TC、batch_size、uptime、loss_time 的公式口径显示"],
            ["多产品多机台", "支持多产品、多机台组和需求计划", "进一步支持单机台 tool_id 粒度"],
            ["Path 相同情景", "可通过 tool_group 聚合能力处理", "补充情景识别，自动走标准公式"],
            ["Path 不同情景", "已有 capacity_matrix 雏形", "新增产品-制程-机台可行性矩阵和分配模型"],
            ["Backup 产能极大化", "LP 可做需求调整", "新增 primary/backup path、启用条件和最大产出目标"],
            ["最佳/最快提升路径", "已有瓶颈和建议措施", "新增改善动作模拟：开放 backup、修复机台、增加 uptime、减少 loss_time"],
        ],
        [4.1, 5.8, 7.3],
    )

    add_heading(doc, "6. 后续优化方式", 1)
    add_heading(doc, "6.1 数据模型优化", 2)
    add_table(
        doc,
        ["新增/调整对象", "建议字段", "作用"],
        [
            ["tool_master", "tool_id, tool_group_id, area, status, uptime, loss_time", "从机台组扩展到单机台能力，支持复杂 path 分配"],
            ["process_master", "process_id, process_name, seq_group", "明确制程维度，支撑制程级约束和瓶颈定位"],
            ["tool_process_capability", "tool_id, product_id, process_id, path_type, can_run, run_time, batch_size", "建立产品-制程-机台可行性矩阵"],
            ["backup_path", "primary_tool_id, backup_tool_id, enable_rule, switch_cost", "表达 backup 可用条件和切换成本"],
            ["scenario_config", "scenario_id, path_same_flag, backup_flag, objective", "自动识别情景并选择标准公式或优化模型"],
        ],
        [3.4, 6.8, 7.0],
    )

    add_heading(doc, "6.2 算法优化", 2)
    add_table(
        doc,
        ["模块", "优化方向", "说明"],
        [
            ["情景识别器", "自动判断情景1-5", "根据产品数量、机台数量、path 是否一致、是否存在 backup 决定计算路径"],
            ["标准公式引擎", "实现公司标准产能公式", "用于 path 相同或单台机情景，输出标准产能、可用小时、损失小时"],
            ["复杂 path 分配模型", "使用 LP/MIP 或最大流替代穷举", "70 台机台、60 个制程下不建议穷举，优化模型更稳定"],
            ["Backup 极大化", "将 backup 作为可选能力或低优先级路径", "目标可设为最大产出、最小切换、最小瓶颈缺口"],
            ["瓶颈解释", "输出约束制程、约束机台、受影响产品", "让业务知道为什么不是总机台数足够就能满足计划"],
            ["提升路径推荐", "对改善动作做边际收益排序", "例如开放某 backup、修复某机台、提升 uptime、降低 loss_time"],
        ],
        [3.4, 5.0, 8.8],
    )

    add_heading(doc, "6.3 前端与导入模板优化", 2)
    add_bullets(
        doc,
        [
            "Excel 模板增加 tool_master、process_master、tool_process_capability、backup_path 等 sheet。",
            "前端导入区提示区分标准情景数据和复杂 path 数据，避免用户只导入机台组产能却期望做复杂 path 优化。",
            "RCCP 热点表增加“约束制程”和“可用机台数”字段，帮助解释复杂 path 下的瓶颈。",
            "生产计划页增加“情景类型”卡片，显示当前走标准公式、分配模型还是 backup 极大化模型。",
            "结果区增加“最佳提升路径”表，展示改善动作、预计提升产能、影响产品、执行难度和推荐优先级。",
        ],
    )

    add_heading(doc, "6.4 验证与上线建议", 2)
    add_table(
        doc,
        ["阶段", "验证方式", "验收标准"],
        [
            ["阶段1：标准公式复现", "用 Excel 中的标准公式案例对齐", "系统产能值与人工公式结果一致"],
            ["阶段2：复杂 path 小样例", "复现两个产品、三个机台、三个制程矩阵", "能正确识别唯一瓶颈和不可分配组合"],
            ["阶段3：Backup 极大化", "构造有/无 backup 的对比案例", "启用 backup 后产出不下降，瓶颈解释合理"],
            ["阶段4：真实 R6 验证", "导入真实 R6 计划、OEE、WIP 和能力矩阵", "能支持正式投片承诺评审和周计划冻结前校验"],
            ["阶段5：结果追溯", "保存数据版本、情景、参数、求解状态和输出结果", "评审留档可复盘，结果可解释"],
        ],
        [3.5, 6.0, 7.7],
    )

    add_heading(doc, "7. 建议实施优先级", 1)
    add_table(
        doc,
        ["优先级", "优化项", "原因"],
        [
            ["P0", "补齐 tool_id、process_id、capability matrix 数据结构", "这是复杂 path 和 backup 建模的基础"],
            ["P0", "新增情景识别器", "先判断走标准公式还是优化模型，避免所有场景混用同一算法"],
            ["P1", "实现标准公式引擎并与公司口径对齐", "快速复现简单情景，提高业务信任"],
            ["P1", "实现 LP/MIP 分配模型", "替代穷举，支撑多机台、多制程、多产品规模化计算"],
            ["P1", "输出最佳提升路径", "贴合源文件第二个核心诉求：给出最佳/最快提升建议"],
            ["P2", "前端增加复杂 path 可视化和导入校验", "让业务能看懂为什么某些机台不能简单合并计算"],
        ],
        [2.2, 5.5, 9.5],
    )

    add_heading(doc, "8. 总结", 1)
    add_body(
        doc,
        "这份 Excel 的核心价值在于定义了公司产能计算的情景分类和业务诉求：简单 path 用标准公式快速计算，复杂 path 和 backup 场景则需要构建产品-制程-机台可行性矩阵，并通过分配模型寻找瓶颈和产能最大化方案。Capacity Agent 当前已经具备 R6、WIP-aware、RCCP、LP 和冻结状态的基础能力；下一步应重点补齐 tool_id/process_id 粒度、复杂 path 分配模型和最佳提升路径推荐，才能更完整地承接公司现有产能计算方法。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
